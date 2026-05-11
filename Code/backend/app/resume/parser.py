"""
Resume parsing logic for different file formats.
Extracts text from PDF, DOCX, and image files.
Uses Tesseract OCR as fallback for scanned PDFs and images.
"""
import pdfplumber
from docx import Document
from typing import Optional
import os
import re

from app.core.exceptions import FileProcessingError


class ResumeParser:
    """Parser for extracting text from resume files.

    Supports:
    - PDF: pdfplumber for native text, Tesseract OCR fallback for scanned PDFs
    - DOCX: python-docx paragraph extraction
    - Images: Tesseract OCR (PNG, JPG, JPEG)
    """

    # Minimum characters to consider a PDF as having extractable text
    MIN_TEXT_LENGTH = 50

    def parse_file(self, file_path: str) -> str:
        """
        Parse resume file and extract text.

        Args:
            file_path: Path to resume file

        Returns:
            Extracted text

        Raises:
            FileProcessingError: If parsing fails
        """
        if not os.path.exists(file_path):
            raise FileProcessingError(f"File not found: {file_path}")

        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        try:
            if ext == '.pdf':
                return self._parse_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            elif ext in ['.png', '.jpg', '.jpeg']:
                return self._parse_image(file_path)
            else:
                raise FileProcessingError(f"Unsupported file format: {ext}")
        except FileProcessingError:
            raise
        except Exception as e:
            raise FileProcessingError(f"Failed to parse file: {str(e)}")

    def _parse_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file.

        Uses pdfplumber first. If the result is too short (likely scanned),
        falls back to Tesseract OCR via pdf2image.
        """
        text = ""

        # Step 1: Native text — prefer longer of default vs layout mode (helps columns)
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                a = page.extract_text() or ""
                try:
                    b = page.extract_text(layout=True) or ""
                except Exception:
                    b = ""
                page_text = b if len(b) > len(a) else a
                if page_text:
                    text += page_text + "\n"

        text = self._clean_text(text)

        # Step 1b: If still thin text, rebuild from word boxes (some templates hide text)
        if len(text) < self.MIN_TEXT_LENGTH:
            text = self._pdf_words_fallback(file_path) or text

        # Step 2: If insufficient text, try OCR fallback (scanned PDF)
        if len(text) < self.MIN_TEXT_LENGTH:
            ocr_text = self._ocr_pdf(file_path)
            if ocr_text and len(ocr_text) > len(text):
                text = ocr_text

        return text

    def _ocr_pdf(self, file_path: str) -> str:
        """
        Extract text from a scanned PDF using Tesseract OCR.

        Converts PDF pages to images, then runs OCR on each page.
        """
        try:
            from pdf2image import convert_from_path
            import pytesseract
            from PIL import ImageFilter

            images = convert_from_path(file_path, dpi=300)
            text = ""

            for img in images:
                # Preprocess: convert to grayscale for better OCR
                img_gray = img.convert('L')
                # Apply slight sharpening
                img_sharp = img_gray.filter(ImageFilter.SHARPEN)

                page_text = pytesseract.image_to_string(img_sharp)
                if page_text:
                    text += page_text + "\n"

            return self._clean_text(text)

        except ImportError as e:
            print(f"OCR PDF fallback unavailable (missing dependency): {e}")
            return ""
        except Exception as e:
            print(f"OCR PDF fallback failed: {e}")
            return ""

    def _pdf_words_fallback(self, file_path: str) -> str:
        """Reconstruct text from word bounding boxes when extract_text under-performs."""
        try:
            lines_out = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    words = page.extract_words(use_text_flow=True)
                    if not words:
                        continue
                    words.sort(key=lambda w: (round(float(w.get("top", 0)) / 3) * 3, float(w.get("x0", 0))))
                    line_parts = []
                    last_top = None
                    for w in words:
                        t = round(float(w.get("top", 0)) / 4) * 4
                        txt = (w.get("text") or "").strip()
                        if not txt:
                            continue
                        if last_top is not None and abs(t - last_top) > 6:
                            if line_parts:
                                lines_out.append(" ".join(line_parts))
                            line_parts = []
                        line_parts.append(txt)
                        last_top = t
                    if line_parts:
                        lines_out.append(" ".join(line_parts))
            return self._clean_text("\n".join(lines_out))
        except Exception as e:
            print(f"PDF words fallback failed: {e}")
            return ""

    def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        lines = []

        # Paragraph text
        for paragraph in doc.paragraphs:
            para = (paragraph.text or "").strip()
            if para:
                lines.append(para)

        # Table text (many resumes place skills/experience in tables)
        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = " ".join(
                        (p.text or "").strip() for p in cell.paragraphs if (p.text or "").strip()
                    ).strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    lines.append(" | ".join(row_cells))

        # Headers & footers (some templates put contact/skills here)
        for sec in doc.sections:
            try:
                for part in (sec.header, sec.footer):
                    if part is None:
                        continue
                    for p in part.paragraphs:
                        t = (p.text or "").strip()
                        if t:
                            lines.append(t)
            except Exception:
                pass

        return self._clean_text("\n".join(lines))

    def _parse_image(self, file_path: str) -> str:
        """
        Extract text from image using Tesseract OCR.

        Applies grayscale conversion and sharpening for better accuracy.
        """
        try:
            import pytesseract
            from PIL import Image, ImageFilter

            image = Image.open(file_path)

            # Preprocess: grayscale + sharpen
            image = image.convert('L')
            image = image.filter(ImageFilter.SHARPEN)

            text = pytesseract.image_to_string(image)

            return self._clean_text(text)
        except ImportError:
            raise FileProcessingError("pytesseract not installed. Install with: pip install pytesseract")
        except Exception as e:
            raise FileProcessingError(f"OCR failed: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text.

        Keep section structure while removing OCR artifacts.
        """
        # Normalize line endings
        text = (text or "").replace("\r\n", "\n").replace("\r", "\n")

        # Keep newlines (important for section extraction), but normalize spacing.
        cleaned_lines = []
        for line in text.split("\n"):
            line = re.sub(r"[ \t]+", " ", line).strip()
            if line:
                cleaned_lines.append(line)

        text = "\n".join(cleaned_lines)

        return text.strip()


# Singleton instance
_parser = None


def get_parser() -> ResumeParser:
    """Get singleton parser instance"""
    global _parser
    if _parser is None:
        _parser = ResumeParser()
    return _parser
